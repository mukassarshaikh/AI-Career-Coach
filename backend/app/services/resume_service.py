"""
resume_service.py — Resume business logic (Phase 1).

Handles:
  - File extension & MIME type validation (PDF, DOCX)
  - Uploading file to Cloudinary storage
  - Text extraction from uploaded PDF / DOCX files (pypdf, python-docx)
  - Database row creation and updates in `resumes` table
  - ATS scoring, grammar auditing, and `resume_reports` row creation
  - Target Job Description creation and keyword gap analysis
  - Fetching resume records and report details by ID for authorized users
"""

import io
import logging
from typing import Optional
from uuid import UUID

import cloudinary
import cloudinary.uploader
import httpx
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.resume import JobDescription, Resume, ResumeReport
from app.services import llm_service

logger = logging.getLogger(__name__)

# Allowed file extensions and MIME types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/octet-stream",  # Fallback for some browsers uploading docx
}


def _configure_cloudinary() -> None:
    """Configures the Cloudinary SDK with environment settings."""
    cloudinary.config(
        cloud_name=settings.cloudinary_cloud_name,
        api_key=settings.cloudinary_api_key,
        api_secret=settings.cloudinary_api_secret,
        secure=True,
    )


async def validate_resume_file(file: UploadFile) -> None:
    """
    Validates that the uploaded file has a valid PDF or DOCX extension, MIME type,
    and authoritative binary magic-byte content signature.
    Raises HTTPException(400) if validation fails.
    """
    filename = file.filename or ""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # 1. First-pass extension filter
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid file extension '{ext}'. Only PDF (.pdf) and Word (.docx) files are allowed.",
        )

    # 2. First-pass Content-Type filter
    if file.content_type and file.content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid file content type '{file.content_type}'. Only PDF and Word documents are accepted.",
        )

    # 3. Authoritative magic-byte binary header validation
    await file.seek(0)
    header_bytes = await file.read(2048)
    await file.seek(0)

    if not header_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file is empty.",
        )

    is_valid = False
    if ext == ".pdf":
        # PDF magic bytes: %PDF- in header
        is_valid = b"%PDF-" in header_bytes[:1024]
    elif ext in (".docx", ".doc"):
        # Word / DOCX magic bytes: PK zip header (PK\x03\x04 or PK\x05\x06) or OLE CFB header (\xd0\xcf\x11...)
        try:
            import filetype

            kind = filetype.guess(header_bytes)
            is_ft_match = kind is not None and kind.extension in ("docx", "doc", "zip")
        except ImportError:
            is_ft_match = False

        is_pk_zip = header_bytes.startswith(b"PK\x03\x04") or header_bytes.startswith(b"PK\x05\x06")
        is_doc_ole = header_bytes.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
        is_valid = is_ft_match or is_pk_zip or is_doc_ole

    if not is_valid:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="This file doesn't appear to be a valid PDF or Word document.",
        )


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extracts raw text from PDF or DOCX file bytes.
    Uses `pypdf` for PDFs and `python-docx` for Word documents.
    Imports are lazy to avoid server startup errors if dependencies are missing.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text_chunks: list[str] = []

    if ext == ".pdf":
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_chunks.append(extracted)
        except ImportError:
            raise ValueError("pypdf package is not installed. Please run 'pip install pypdf'.")
        except Exception as exc:
            raise ValueError(f"Failed to extract text from PDF: {exc}") from exc
    elif ext in (".docx", ".doc"):
        try:
            import docx

            doc = docx.Document(io.BytesIO(file_bytes))
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_chunks.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_chunks.append(" | ".join(row_text))
        except ImportError:
            raise ValueError("python-docx package is not installed. Please run 'pip install python-docx'.")
        except Exception as exc:
            raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc
    else:
        # Fallback text decoding
        return file_bytes.decode("utf-8", errors="ignore")

    return "\n\n".join(text_chunks).strip()


def get_signed_resume_url(public_id: str) -> str:
    """
    Generates a short-lived (~1 hour) signed delivery URL for an authenticated
    Cloudinary raw resume asset using Cloudinary SDK conventions.
    """
    _configure_cloudinary()
    url, _ = cloudinary.utils.cloudinary_url(
        public_id,
        resource_type="raw",
        type="authenticated",
        sign_url=True,
        secure=True,
    )
    return url


def delete_cloudinary_asset(
    public_id: str,
    resource_type: str = "raw",
    delivery_type: str = "authenticated",
) -> dict:
    """
    Deletes a Cloudinary asset by public_id and delivery_type.
    Distinguishes:
      - successful deletion (result='ok') -> returns {"status": "ok", "result": "ok"}
      - asset already absent (result='not found') -> returns {"status": "ok", "result": "not_found"}
      - unconfigured/mock env -> returns {"status": "ok", "result": "mock_skipped"}
      - genuine failure -> returns {"status": "failed", "error": "..."}
    """
    if not public_id:
        return {"status": "ok", "result": "not_applicable"}

    if not settings.cloudinary_cloud_name or "your_cloud_name" in settings.cloudinary_cloud_name:
        logger.info(f"Cloudinary unconfigured/mock env; skipping physical destruction for '{public_id}'")
        return {"status": "ok", "result": "mock_skipped"}

    try:
        _configure_cloudinary()
        res = cloudinary.uploader.destroy(public_id, resource_type=resource_type, type=delivery_type)

        res_str = str(res).lower()
        result_val = res.get("result") if isinstance(res, dict) else str(res)
        result_str = str(result_val).lower()

        if result_str == "ok":
            logger.info(f"Cloudinary destruction for '{public_id}' ({delivery_type}) succeeded: result='ok'")
            return {"status": "ok", "result": "ok"}
        elif "not found" in res_str or "404" in res_str or result_str in ("not found", "not_found"):
            logger.info(f"Cloudinary destruction for '{public_id}' ({delivery_type}) returned not found ({res}); treating as already erased.")
            return {"status": "ok", "result": "not found"}
        else:
            logger.error(f"Cloudinary destruction for '{public_id}' ({delivery_type}) failed: {res}")
            return {"status": "failed", "error": f"Cloudinary response: {res}"}
    except Exception as exc:
        exc_str = str(exc).lower()
        if "not found" in exc_str or "404" in exc_str:
            logger.info(f"Cloudinary asset '{public_id}' ({delivery_type}) not found via exception ({exc}); treating as already erased.")
            return {"status": "ok", "result": "not found"}
        logger.error(f"Exception destroying Cloudinary asset '{public_id}': {exc}")
        return {"status": "failed", "error": str(exc)}


def extract_cloudinary_info_from_resume(resume: Resume) -> tuple[Optional[str], str]:
    """
    Extracts (public_id, delivery_type) for a resume record.
    - If `cloudinary_public_id` is set (authenticated uploads), returns (cloudinary_public_id, "authenticated").
    - If `cloudinary_public_id` is None (legacy uploads), parses `file_url` if it is a Cloudinary URL.
    - Returns (None, "authenticated") if `file_url` is not a Cloudinary URL.
    """
    if resume.cloudinary_public_id:
        return resume.cloudinary_public_id, "authenticated"

    file_url = resume.file_url or ""
    if "cloudinary.com" in file_url:
        delivery_type = "authenticated" if "/raw/authenticated/" in file_url else "upload"
        marker = "/raw/authenticated/" if "/raw/authenticated/" in file_url else "/raw/upload/"
        if marker in file_url:
            path_part = file_url.split(marker, 1)[1]
            import re
            clean_path = re.sub(r"^v\d+/", "", path_part)
            public_id = clean_path.split("?")[0]
            return public_id, delivery_type

    return None, "authenticated"


def get_effective_resume_file_url(resume: Resume) -> str:
    """
    Returns a valid delivery URL for the resume:
    - If `cloudinary_public_id` is set (new authenticated uploads), generates a fresh signed URL.
    - If `cloudinary_public_id` is None (legacy rows), falls back to stored `file_url`.
    """
    if resume.cloudinary_public_id:
        return get_signed_resume_url(resume.cloudinary_public_id)
    return resume.file_url


async def extract_text_from_url(file_url: str) -> str:
    """
    Downloads file bytes from a Cloudinary URL and extracts raw text.
    """
    async with httpx.AsyncClient() as client:
        response = await client.get(file_url, follow_redirects=True, timeout=30.0)
        response.raise_for_status()

    filename = file_url.split("?")[0].split("/")[-1]
    return extract_text_from_bytes(response.content, filename)


async def upload_resume_file(
    file: UploadFile,
    user_id: UUID,
    db: AsyncSession,
) -> Resume:
    """
    Validates the file, uploads it to Cloudinary as an authenticated private raw asset,
    stores the permanent `cloudinary_public_id`, and saves a `resumes` record in the database.
    """
    # 1. Validate file extension and format (including binary magic bytes)
    await validate_resume_file(file)

    # 2. Upload to Cloudinary using authenticated delivery mode
    _configure_cloudinary()
    try:
        await file.seek(0)
        file_bytes = await file.read()

        public_id_str = f"resumes/user_{user_id}_{file.filename}"
        upload_result = cloudinary.uploader.upload(
            file_bytes,
            resource_type="raw",
            type="authenticated",
            public_id=public_id_str,
            overwrite=True,
        )
        returned_public_id = upload_result.get("public_id") or public_id_str
        file_url = (
            upload_result.get("secure_url")
            or upload_result.get("url")
            or f"https://res.cloudinary.com/{settings.cloudinary_cloud_name}/raw/authenticated/{returned_public_id}"
        )
        if not file_url:
            raise ValueError("Cloudinary response did not return a valid file URL.")
    except Exception as exc:
        if "your_cloud_name" in settings.cloudinary_cloud_name or not settings.cloudinary_cloud_name:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Cloudinary is not configured. Please set valid Cloudinary credentials in .env.",
            ) from exc
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload resume file to Cloudinary: {str(exc)}",
        ) from exc

    # 3. Create database row with permanent cloudinary_public_id
    resume = Resume(
        user_id=user_id,
        file_url=file_url,
        cloudinary_public_id=returned_public_id,
        raw_text=None,
        parsed_json=None,
        ats_score=None,
    )
    db.add(resume)
    await db.commit()
    await db.refresh(resume)

    return resume


async def get_resume_by_id(
    db: AsyncSession,
    resume_id: UUID,
    user_id: Optional[UUID] = None,
) -> Optional[Resume]:
    """
    Fetches a resume by ID (optionally matching user_id).
    """
    stmt = select(Resume).where(Resume.id == resume_id)
    if user_id is not None:
        stmt = stmt.where(Resume.user_id == user_id)
    result = await db.execute(stmt)
    return result.scalar_one_or_none()


async def process_parse_resume_job(
    db: AsyncSession,
    resume_id: UUID,
) -> dict:
    """
    Core shared business logic for parsing a resume record by ID:
      1. Fetches the `resumes` row from Postgres by `resume_id`.
      2. Downloads document from Cloudinary URL and extracts raw text.
      3. Calls Groq LLM (`llm_service.structure_resume`) to structure raw text into JSON.
      4. Logs the AI call to `ai_generation_logs` (module='resume').
      5. Updates `resumes.raw_text` and `resumes.parsed_json` in Postgres.
    """
    logger.info(f"Starting process_parse_resume_job for resume_id={resume_id}")

    resume = await get_resume_by_id(db, resume_id=resume_id)
    if not resume:
        logger.error(f"Resume {resume_id} not found in database.")
        raise ValueError(f"Resume {resume_id} not found in database.")

    # 1. Extract text from Cloudinary file URL
    effective_url = get_effective_resume_file_url(resume)
    raw_text = await extract_text_from_url(effective_url)

    # 2. Call Groq LLM to structure resume data into JSON
    parsed_json = await llm_service.structure_resume(
        text=raw_text,
        user_id=resume.user_id,
        db=db,
    )

    # 3. Save results to resumes table
    resume.raw_text = raw_text
    resume.parsed_json = parsed_json
    await db.commit()

    logger.info(f"Completed process_parse_resume_job for resume_id={resume_id}")
    return {
        "status": "complete",
        "resume_id": str(resume_id),
        "raw_text_length": len(raw_text),
        "parsed": True,
    }


async def list_user_resumes(
    db: AsyncSession,
    user_id: UUID,
) -> list[Resume]:
    """
    Fetches all Resume records owned by user_id, ordered by created_at descending.
    """
    stmt = select(Resume).where(Resume.user_id == user_id).order_by(Resume.created_at.desc())
    result = await db.execute(stmt)
    return list(result.scalars().all())


async def create_resume_report(
    db: AsyncSession,
    resume: Resume,
) -> ResumeReport:
    """
    Evaluates ATS score and performs grammar audit on a parsed resume.
    Updates `resumes.ats_score` and inserts a `resume_reports` database row.
    Leaves `keyword_gaps`, `action_items`, and `job_description_id` as null.
    """
    if not resume.parsed_json:
        raise ValueError("Cannot score resume: parsed_json is null. Run parse_resume job first.")

    raw_text = resume.raw_text or ""

    # 1. Call LLM for ATS scoring
    ats_data = await llm_service.score_resume_ats(
        parsed_json=resume.parsed_json,
        raw_text=raw_text,
        user_id=resume.user_id,
        db=db,
    )

    # 2. Call LLM for grammar audit
    grammar_data = await llm_service.audit_resume_grammar(
        raw_text=raw_text,
        user_id=resume.user_id,
        db=db,
    )

    # 3. Update overall ATS score on the resume record
    overall_score = ats_data.get("overall_score", 70)
    try:
        resume.ats_score = int(overall_score)
    except (ValueError, TypeError):
        resume.ats_score = 70

    # 4. Insert resume_reports row matching database.md schema
    report = ResumeReport(
        resume_id=resume.id,
        job_description_id=None,
        ats_breakdown=ats_data,
        grammar_suggestions=grammar_data.get("suggestions", []),
        keyword_gaps=None,
        action_items=None,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)

    return report


async def get_latest_resume_report(
    db: AsyncSession,
    resume_id: UUID,
    user_id: UUID,
) -> Optional[ResumeReport]:
    """
    Fetches the most recent ResumeReport for a resume belonging to the authenticated user.
    """
    stmt = (
        select(ResumeReport)
        .join(Resume, Resume.id == ResumeReport.resume_id)
        .where(ResumeReport.resume_id == resume_id, Resume.user_id == user_id)
        .order_by(ResumeReport.created_at.desc())
        .limit(1)
    )
    result = await db.execute(stmt)
    return result.scalar_one_or_none()


async def create_job_description(
    db: AsyncSession,
    user_id: UUID,
    resume_id: UUID,
    raw_text: str,
) -> JobDescription:
    """
    Inserts a new target JobDescription record into Postgres linked to user_id and resume_id.
    """
    jd = JobDescription(
        user_id=user_id,
        resume_id=resume_id,
        raw_text=raw_text,
        parsed_keywords=None,
    )
    db.add(jd)
    await db.commit()
    await db.refresh(jd)
    return jd


async def get_job_description_by_id(
    db: AsyncSession,
    job_description_id: UUID,
) -> Optional[JobDescription]:
    """
    Fetches a JobDescription by ID.
    """
    stmt = select(JobDescription).where(JobDescription.id == job_description_id)
    result = await db.execute(stmt)
    return result.scalar_one_or_none()


async def analyze_resume_keywords(
    db: AsyncSession,
    job_description_id: UUID,
    resume_id: UUID,
) -> ResumeReport:
    """
    Compares a candidate's resume against a target job description:
      1. Fetches resume and target job_description.
      2. Fetches the resume's most recent prior `ResumeReport` row.
         - Fails cleanly if no prior scoring report exists (must score resume first).
      3. Calls Groq `analyze_keywords_llm` to extract matched keywords, missing keywords, and action items.
      4. Updates `job_description.parsed_keywords`.
      5. Creates a NEW `ResumeReport` row carrying forward `ats_breakdown` and `grammar_suggestions`
         from the prior report, with `job_description_id`, `keyword_gaps`, and `action_items` populated.
    """
    resume = await get_resume_by_id(db, resume_id=resume_id)
    if not resume:
        raise ValueError(f"Resume {resume_id} not found.")

    jd = await get_job_description_by_id(db, job_description_id=job_description_id)
    if not jd:
        raise ValueError(f"Job description {job_description_id} not found.")

    # Fetch most recent prior ResumeReport to carry forward ats_breakdown and grammar_suggestions
    prior_report = await get_latest_resume_report(db, resume_id=resume_id, user_id=resume.user_id)
    if not prior_report:
        raise ValueError(
            f"Cannot analyze keywords for resume {resume_id}: no prior scoring report found. Run score_resume job first."
        )

    # Call LLM for keyword gap analysis
    analysis_result = await llm_service.analyze_keywords_llm(
        resume_text=resume.raw_text or "",
        jd_text=jd.raw_text,
        user_id=resume.user_id,
        db=db,
    )

    # Save parsed keywords summary to JobDescription
    jd.parsed_keywords = {
        "matched": analysis_result.get("matched_keywords", []),
        "missing": analysis_result.get("missing_keywords", []),
    }

    # DATA MODELING DECISION: Create a NEW ResumeReport row carrying forward prior ATS/grammar data
    new_report = ResumeReport(
        resume_id=resume.id,
        job_description_id=jd.id,
        ats_breakdown=prior_report.ats_breakdown,
        grammar_suggestions=prior_report.grammar_suggestions,
        keyword_gaps=analysis_result.get("missing_keywords", []),
        action_items=analysis_result.get("action_items", []),
    )
    db.add(new_report)
    await db.commit()
    await db.refresh(new_report)

    return new_report
