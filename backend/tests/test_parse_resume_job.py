"""
Pytest integration test suite for parse_resume Arq job, text extraction, LLM structuring, and job polling.

Run with:
    python -m pytest tests/test_parse_resume_job.py -v
"""

import io
import json
import uuid
from datetime import datetime
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import UUID

import docx
import pypdf
import pytest
from fastapi.testclient import TestClient
from groq import RateLimitError
from jose import jwt

from app.core.config import settings
from app.core.security import ALGORITHM
from app.main import app
from app.models.logs import AiGenerationLog
from app.models.resume import Resume
from app.models.user import User
from app.services import llm_service, resume_service
from app.workers.jobs.parse_resume import parse_resume

client = TestClient(app)


def generate_test_token(email: str = "testuser@example.com") -> str:
    """Generates a valid NextAuth JWT Bearer token for authentication."""
    payload = {"sub": email, "email": email, "name": "Test User"}
    return jwt.encode(payload, settings.nextauth_secret, algorithm=ALGORITHM)


# ---------------------------------------------------------------------------
# 1. Text Extraction Tests (pypdf & python-docx)
# ---------------------------------------------------------------------------
def test_extract_text_from_pdf():
    """Verify text extraction from in-memory PDF file bytes using pypdf."""
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=100, height=100)
    pdf_buffer = io.BytesIO()
    writer.write(pdf_buffer)
    pdf_bytes = pdf_buffer.getvalue()

    extracted = resume_service.extract_text_from_bytes(pdf_bytes, "my_resume.pdf")
    assert isinstance(extracted, str)


def test_extract_text_from_docx():
    """Verify text extraction from in-memory DOCX file bytes using python-docx."""
    doc = docx.Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Software Engineer with 5 years experience in Python and FastAPI.")
    doc.add_paragraph("Education: BS Computer Science")

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    extracted = resume_service.extract_text_from_bytes(docx_bytes, "my_resume.docx")
    assert "Jane Doe" in extracted
    assert "Software Engineer" in extracted
    assert "BS Computer Science" in extracted


# ---------------------------------------------------------------------------
# 2. LLM Structuring & Audit Logging Tests (Groq)
# ---------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_structure_resume_llm_success():
    """Verify llm_service.structure_resume parses Groq output and logs to ai_generation_logs."""
    dummy_llm_json = {
        "experience": [
            {
                "company": "Tech Corp",
                "role": "Senior Engineer",
                "start_date": "2021-01",
                "end_date": "Present",
                "description": "Led backend team",
                "highlights": ["Built FastAPI microservices"],
            }
        ],
        "education": [
            {
                "institution": "State University",
                "degree": "BS CS",
                "field_of_study": "Computer Science",
                "graduation_year": "2020",
            }
        ],
        "skills": {
            "technical": ["Python", "FastAPI", "Postgres"],
            "tools": ["Docker", "Git"],
            "soft_skills": ["Leadership"],
        },
        "achievements": ["Top Performer 2023"],
    }

    mock_chat_completion = MagicMock()
    mock_chat_completion.choices = [
        MagicMock(message=MagicMock(content=json.dumps(dummy_llm_json)))
    ]

    mock_db = AsyncMock()
    mock_db.add = MagicMock()

    mock_groq_client = MagicMock()
    mock_groq_client.chat.completions.create = AsyncMock(return_value=mock_chat_completion)

    with patch("app.services.llm_service._get_groq_client", return_value=mock_groq_client):
        result = await llm_service.structure_resume(
            text="Resume content string",
            user_id=uuid.uuid4(),
            db=mock_db,
        )

        assert result["experience"][0]["company"] == "Tech Corp"
        assert "Python" in result["skills"]["technical"]

        # Verify audit log entry was created in database session
        assert mock_db.add.called
        saved_log = mock_db.add.call_args[0][0]
        assert isinstance(saved_log, AiGenerationLog)
        assert saved_log.module == "resume"
        assert saved_log.model_used == settings.groq_model


@pytest.mark.asyncio
async def test_structure_resume_groq_rate_limit_retry():
    """Verify structure_resume handles 429 rate limit errors with backoff retry."""
    dummy_llm_json = {"experience": [], "education": [], "skills": {}, "achievements": []}

    mock_chat_completion = MagicMock()
    mock_chat_completion.choices = [
        MagicMock(message=MagicMock(content=json.dumps(dummy_llm_json)))
    ]

    # Create dummy 429 RateLimitError
    dummy_response = MagicMock(status_code=429, headers={})
    rate_limit_err = RateLimitError(
        message="Rate limit exceeded",
        response=dummy_response,
        body={"error": {"message": "Rate limit exceeded"}},
    )

    mock_db = AsyncMock()
    mock_db.add = MagicMock()

    mock_groq_client = MagicMock()
    mock_groq_client.chat.completions.create = AsyncMock(
        side_effect=[rate_limit_err, rate_limit_err, mock_chat_completion]
    )

    with patch("app.services.llm_service._get_groq_client", return_value=mock_groq_client), \
         patch("asyncio.sleep", new_callable=AsyncMock) as mock_sleep:

        result = await llm_service.structure_resume(
            text="Test text",
            user_id=uuid.uuid4(),
            db=mock_db,
            max_retries=3,
        )

        assert result == dummy_llm_json
        assert mock_groq_client.chat.completions.create.call_count == 3
        assert mock_sleep.call_count == 2


# ---------------------------------------------------------------------------
# 3. Arq Worker Job Execution Test
# ---------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_parse_resume_worker_job():
    """Verify parse_resume Arq worker job extracts text, structures JSON, and updates resumes DB row."""
    resume_id = str(uuid.uuid4())
    user_id = uuid.uuid4()

    mock_resume = Resume(
        id=UUID(resume_id),
        user_id=user_id,
        file_url="https://res.cloudinary.com/demo/raw/upload/resumes/resume.pdf",
        raw_text=None,
        parsed_json=None,
    )

    mock_db = AsyncMock()

    class DummyAsyncSessionContext:
        async def __aenter__(self):
            return mock_db

        async def __aexit__(self, exc_type, exc_val, exc_tb):
            pass

    ctx = {"db_factory": lambda: DummyAsyncSessionContext()}

    structured_data = {
        "experience": [{"company": "Acme", "role": "Developer"}],
        "education": [],
        "skills": {"technical": ["Python"]},
        "achievements": [],
    }

    with patch("app.services.resume_service.get_resume_by_id", AsyncMock(return_value=mock_resume)), \
         patch("app.services.resume_service.extract_text_from_url", AsyncMock(return_value="Extracted resume text content")), \
         patch("app.services.llm_service.structure_resume", AsyncMock(return_value=structured_data)):

        result = await parse_resume(ctx, resume_id)

        assert result["status"] == "complete"
        assert result["resume_id"] == resume_id
        assert mock_resume.raw_text == "Extracted resume text content"
        assert mock_resume.parsed_json == structured_data
        assert mock_db.commit.called


# ---------------------------------------------------------------------------
# 4. API Endpoints & Polling Route Tests
# ---------------------------------------------------------------------------
def test_upload_enqueues_parse_job_and_returns_job_id():
    """Verify POST /api/v1/resume/upload enqueues parse_resume and returns job_id in response."""
    token = generate_test_token()
    headers = {"Authorization": f"Bearer {token}"}

    dummy_user = User(id=uuid.uuid4(), email="testuser@example.com", name="Test User")
    dummy_resume = Resume(
        id=uuid.uuid4(),
        user_id=dummy_user.id,
        file_url="https://res.cloudinary.com/demo/raw/upload/resumes/test.pdf",
        created_at=datetime.now(),
    )

    from app.api.v1.deps import get_current_user, get_db
    app.dependency_overrides[get_current_user] = lambda: dummy_user

    mock_db = AsyncMock()
    async def mock_get_db_gen():
        yield mock_db

    app.dependency_overrides[get_db] = mock_get_db_gen

    try:
        with patch("app.services.resume_service.upload_resume_file", AsyncMock(return_value=dummy_resume)), \
             patch("app.api.v1.resume._enqueue_parse_job", AsyncMock(return_value="job_test_12345")):

            files = {"file": ("resume.pdf", b"%PDF-1.4 dummy", "application/pdf")}
            response = client.post("/api/v1/resume/upload", files=files, headers=headers)

            assert response.status_code == 201
            data = response.json()
            assert data["job_id"] == "job_test_12345"
            assert data["resume_id"] == str(dummy_resume.id)
    finally:
        app.dependency_overrides.clear()


def test_get_job_status_polling_route():
    """Verify GET /api/v1/resume/jobs/{job_id} returns JobStatusResponse schema for polling."""
    token = generate_test_token()
    headers = {"Authorization": f"Bearer {token}"}

    dummy_user = User(id=uuid.uuid4(), email="testuser@example.com")
    from app.api.v1.deps import get_current_user
    app.dependency_overrides[get_current_user] = lambda: dummy_user

    try:
        response = client.get("/api/v1/resume/jobs/job_test_12345", headers=headers)
        assert response.status_code == 200
        data = response.json()
        assert "job_id" in data
        assert data["job_id"] == "job_test_12345"
        assert data["status"] in ["queued", "in_progress", "complete", "failed"]
    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_parse_resume_uses_signed_url_for_authenticated_resumes():
    """
    Verify parse_resume job worker passes the signed URL generated by
    get_effective_resume_file_url to extract_text_from_url for new authenticated resumes.
    """
    resume_id = str(uuid.uuid4())
    mock_resume = Resume(
        id=UUID(resume_id),
        user_id=uuid.uuid4(),
        file_url="https://res.cloudinary.com/demo/raw/authenticated/resumes/auth.pdf",
        cloudinary_public_id="resumes/auth.pdf",
    )

    mock_db = AsyncMock()
    mock_db.commit = AsyncMock()

    class DummyAsyncSessionContext:
        async def __aenter__(self):
            return mock_db
        async def __aexit__(self, exc_type, exc_val, exc_tb):
            pass

    ctx = {"db_factory": lambda: DummyAsyncSessionContext()}

    with patch("app.services.resume_service.get_resume_by_id", AsyncMock(return_value=mock_resume)), \
         patch("app.services.resume_service.get_signed_resume_url", return_value="https://res.cloudinary.com/demo/raw/authenticated/s--signed_token--/resumes/auth.pdf") as mock_sign, \
         patch("app.services.resume_service.extract_text_from_url", AsyncMock(return_value="Extracted text from signed URL")) as mock_extract, \
         patch("app.services.llm_service.structure_resume", AsyncMock(return_value={"skills": {}})):

        result = await parse_resume(ctx, resume_id)

        assert result["status"] == "complete"
        # Assert signed URL was generated from public_id
        mock_sign.assert_called_once_with("resumes/auth.pdf")
        # Assert extract_text_from_url received the signed URL
        mock_extract.assert_called_once_with("https://res.cloudinary.com/demo/raw/authenticated/s--signed_token--/resumes/auth.pdf")

